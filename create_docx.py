#!/usr/bin/env python3
"""
Convert APJ Summit Plan from Markdown to DOCX
============================================

This script converts the APJ_SUMMIT_PLAN.md file to a professional DOCX document
with proper formatting, tables, and styling.

Author: Snowflake Professional Services
Date: October 2025
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def add_table_style(document):
    """Add a custom table style to the document"""
    try:
        # Try to get existing style first
        table_style = document.styles['APJ Summit Table']
    except:
        # Create new table style
        table_style = document.styles.add_style('APJ Summit Table', WD_STYLE_TYPE.TABLE)
        table_style.base_style = document.styles['Table Grid']
    
    return table_style


def create_styled_document():
    """Create a new document with custom styles"""
    doc = Document()
    
    # Modify existing styles
    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Arial'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    
    # Heading 3 style
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.name = 'Arial'
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    return doc


def parse_table_from_markdown(lines):
    """Parse a markdown table and return headers and rows"""
    headers = []
    rows = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
            
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        
        if i == 0:
            headers = cells
        elif '---' not in line:  # Skip separator line
            rows.append(cells)
    
    return headers, rows


def add_table_to_doc(doc, headers, rows):
    """Add a formatted table to the document"""
    if not headers or not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = header.strip('*')  # Remove markdown bold
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                # Clean up cell data
                clean_data = cell_data.strip('*`')  # Remove markdown formatting
                clean_data = re.sub(r'^[✅❌📊🥉🥈🥇]\s*', '', clean_data)  # Remove emojis
                row_cells[i].text = clean_data
    
    # Auto-fit table
    table.autofit = True
    
    return table


def clean_text(text):
    """Clean markdown formatting from text"""
    # Remove markdown formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'__(.*?)__', r'\1', text)      # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)        # Code
    text = re.sub(r'^#+\s*', '', text)            # Remove heading markers
    text = re.sub(r'^>\s*', '', text)             # Remove blockquote markers
    text = re.sub(r'^-\s*', '• ', text)           # Convert list items
    
    # Remove emojis from start of line but keep some key ones
    text = re.sub(r'^[🎯🏗️📊🚀📋🧪🏢📈🤖🧠⏱️🛠️🔄]\s*', '', text)
    
    return text.strip()


def convert_markdown_to_docx(markdown_file, docx_file):
    """Convert markdown file to DOCX with proper formatting"""
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create styled document
    doc = create_styled_document()
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    table_lines = []
    in_table = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and separators
        if not line or line == '---':
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            # Add code text with monospace formatting
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(line)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table, process it
            headers, rows = parse_table_from_markdown(table_lines)
            add_table_to_doc(doc, headers, rows)
            doc.add_paragraph()  # Add space after table
            in_table = False
            table_lines = []
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = clean_text(line)
            
            if level == 1:
                if 'APJ Tech Summit' in heading_text:
                    doc.add_heading(heading_text, level=0)  # Title
                else:
                    doc.add_heading(heading_text, level=1)
            elif level == 2:
                doc.add_heading(heading_text, level=1)
            elif level == 3:
                doc.add_heading(heading_text, level=2)
            elif level == 4:
                doc.add_heading(heading_text, level=3)
        
        # Handle blockquotes (mission statements)
        elif line.startswith('>'):
            quote_text = clean_text(line)
            quote_para = doc.add_paragraph(quote_text)
            quote_para.italic = True
            quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle list items
        elif line.startswith('-') or line.startswith('*'):
            list_text = clean_text(line)
            doc.add_paragraph(list_text, style='List Bullet')
        
        # Handle regular paragraphs
        else:
            clean_line = clean_text(line)
            if clean_line:
                # Check for special formatting
                if '**' in line:
                    para = doc.add_paragraph()
                    # Split by bold markers and add runs
                    parts = re.split(r'\*\*(.*?)\*\*', line)
                    for j, part in enumerate(parts):
                        run = para.add_run(clean_text(part))
                        if j % 2 == 1:  # Bold parts
                            run.font.bold = True
                else:
                    doc.add_paragraph(clean_line)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        headers, rows = parse_table_from_markdown(table_lines)
        add_table_to_doc(doc, headers, rows)
    
    # Document formatting completed
    
    # Save document
    doc.save(docx_file)
    print(f"✅ DOCX document created successfully: {docx_file}")


def main():
    """Main execution function"""
    
    markdown_file = 'Documentation/APJ_SUMMIT_PLAN.md'
    docx_file = 'Documentation/APJ_SUMMIT_PLAN.docx'
    
    try:
        print("🔄 Converting Markdown to DOCX...")
        convert_markdown_to_docx(markdown_file, docx_file)
        
        print(f"\n📊 Document Statistics:")
        print(f"   Source: {markdown_file}")
        print(f"   Output: {docx_file}")
        print(f"   Status: ✅ Conversion completed successfully")
        
        # Also create the quick reference DOCX
        quick_ref_md = 'Documentation/APJ_SUMMIT_QUICK_REFERENCE.md'
        quick_ref_docx = 'Documentation/APJ_SUMMIT_QUICK_REFERENCE.docx'
        
        if os.path.exists(quick_ref_md):
            print(f"\n🔄 Converting Quick Reference to DOCX...")
            convert_markdown_to_docx(quick_ref_md, quick_ref_docx)
            print(f"   Output: {quick_ref_docx}")
            print(f"   Status: ✅ Quick Reference conversion completed")
        
    except Exception as e:
        print(f"❌ Error converting to DOCX: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
